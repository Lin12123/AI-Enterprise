"""文档正文抽取：把 PDF / Word 转成纯文本，供 LLM 理解并抽取规则。

说明
----
- 云平台(cloud/)独立 venv，可用 pdfplumber / python-docx；两者均在函数内延迟 import，
  即使本机未安装 cloud venv 依赖，也不影响其它端点导入本模块。
- 图片(png/jpg)本轮暂缓 OCR，抽取返回空串并附提示，由上层落草稿附件即可。
- 抽取结果做长度上限截断，避免超长文本撑爆 LLM 上下文。
"""
from __future__ import annotations

import os


# 传给 LLM 的正文最大字符数(过长会截断，防止上下文溢出)
MAX_TEXT_CHARS = 24000


class DocExtractError(RuntimeError):
    """文档解析失败(依赖缺失/文件损坏)时抛出。"""


def extract_text(fmt: str, abs_path: str) -> str:
    """按扩展名抽取文档正文纯文本。

    - pdf  -> pdfplumber 逐页 extract_text 拼接
    - docx -> python-docx 段落 + 表格单元格
    - png/jpg/jpeg -> 暂不支持(返回空串)
    未知格式返回空串。抽取结果超过 MAX_TEXT_CHARS 会截断。
    """
    fmt = (fmt or "").lower().lstrip(".")
    if fmt == "pdf":
        text = _extract_pdf(abs_path)
    elif fmt in ("docx", "doc"):
        text = _extract_docx(abs_path)
    elif fmt in ("png", "jpg", "jpeg", "bmp", "tif", "tiff"):
        # 图片 OCR 暂缓
        return ""
    else:
        return ""

    text = (text or "").strip()
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
    return text


def is_text_extractable(fmt: str) -> bool:
    """该格式当前是否支持自动抽取正文文本(图片暂不支持)。"""
    return (fmt or "").lower().lstrip(".") in ("pdf", "docx", "doc")


def _extract_pdf(abs_path: str) -> str:
    try:
        import pdfplumber  # 延迟 import
    except ImportError as exc:  # pragma: no cover - 依赖缺失路径
        raise DocExtractError("未安装 pdfplumber，无法解析 PDF") from exc

    if not os.path.exists(abs_path):
        raise DocExtractError(f"文件不存在: {abs_path}")

    parts: list[str] = []
    try:
        with pdfplumber.open(abs_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(page_text)
    except Exception as exc:  # pragma: no cover
        raise DocExtractError(f"PDF 解析失败: {exc}") from exc
    return "\n".join(parts)


def _extract_docx(abs_path: str) -> str:
    try:
        import docx  # python-docx，延迟 import
    except ImportError as exc:  # pragma: no cover
        raise DocExtractError("未安装 python-docx，无法解析 Word") from exc

    if not os.path.exists(abs_path):
        raise DocExtractError(f"文件不存在: {abs_path}")

    parts: list[str] = []
    try:
        document = docx.Document(abs_path)
        for para in document.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())
        # 表格：常见于标准中的参数表，逐行按“ | ”拼接
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:  # pragma: no cover
        raise DocExtractError(f"Word 解析失败: {exc}") from exc
    return "\n".join(parts)